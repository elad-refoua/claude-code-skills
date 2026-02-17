"""
Citation Context Extractor: Extract citation-sentence pairs from an academic paper.

Usage: py ref_context.py <input.docx>
Output: <input>_PAIRS.json in the same folder

Extracts each citation-sentence pair with its full reference text.
The JSON output is then used by Claude Code's Sonnet sub-agent for
web search + verification.

Requires: python-docx (py -m pip install python-docx)
No API keys needed - this script only does extraction.
"""

import re
import sys
import json
from pathlib import Path
from docx import Document


def normalize_text(text):
    """Normalize unicode quotes and hyphens."""
    text = text.replace('\u2018', "'").replace('\u2019', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2010', '-').replace('\u2011', '-')
    text = text.replace('\u2012', '-').replace('\u2013', '-')
    text = text.replace('\u200f', '').replace('\u200e', '').replace('\u200b', '')
    return text


def find_references_heading_index(doc):
    """Find the paragraph index where 'References' heading starts."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in ('References', 'REFERENCES', 'Reference List', 'Bibliography'):
            return i
    return len(doc.paragraphs)


def extract_sentences_with_citations(body_text):
    """
    Extract sentences that contain citations, along with the citation info.
    Returns list of dicts: {sentence, citations: [(author, year), ...], paragraph_num}
    """
    text = normalize_text(body_text)
    paragraphs = text.split('\n\n')
    results = []

    for para_num, para in enumerate(paragraphs):
        if not para.strip():
            continue

        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', para)

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            citations_in_sentence = []

            # Parenthetical: (Author, Year)
            for m in re.finditer(r'\(([^)]*\d{4}[a-z]?[^)]*)\)', sentence):
                block = m.group(1)
                for part in block.split(';'):
                    part = part.strip()
                    year_match = re.findall(r'(\d{4}[a-z]?)', part)
                    if year_match:
                        author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                        author_part = re.sub(
                            r'^(e\.g\.,?\s*|see\s+|cf\.\s*|also\s+)',
                            '', author_part, flags=re.IGNORECASE
                        ).strip().rstrip(',')
                        if author_part and len(author_part) > 2:
                            for y in year_match:
                                citations_in_sentence.append((author_part, y))

            # Narrative: Author (Year)
            for m in re.finditer(
                r'([A-Z][a-z]+(?:-[A-Z][a-z]+)?(?:\s+et\s+al\.?)?)\s*\((\d{4}[a-z]?)\)',
                sentence
            ):
                citations_in_sentence.append((m.group(1), m.group(2)))

            # Bracket: Author [Year]
            for m in re.finditer(
                r'([A-Z][a-z]+(?:-[A-Z][a-z]+)?)\s*\[(\d{4}[a-z]?)\]',
                sentence
            ):
                citations_in_sentence.append((m.group(1), m.group(2)))

            if citations_in_sentence:
                unique_cites = list(set(citations_in_sentence))
                results.append({
                    'sentence': sentence[:500],
                    'citations': unique_cites,
                    'paragraph_num': para_num
                })

    return results


def extract_reference_entries(doc, ref_para_idx):
    """Extract full reference text entries as a dict: (surname, year) -> full text."""
    refs = {}
    for pi in range(ref_para_idx + 1, len(doc.paragraphs)):
        text = normalize_text(doc.paragraphs[pi].text.strip())
        if not text:
            continue
        if text in ('References', 'REFERENCES', 'Reference List', 'Bibliography'):
            continue

        year_match = re.search(r'\((\d{4}[a-z]?)\)', text)
        if year_match:
            year = year_match.group(1)
            before = text[:text.index('(')].strip().rstrip(',').strip()
            parts = before.split(',')[0].strip().rstrip('.').split()
            surname = parts[0] if parts else ''
            if surname and len(surname) > 1:
                refs[(surname, year)] = text

    return refs


def build_citation_ref_pairs(sentences_with_cites, ref_entries):
    """
    Match each sentence's citations to their reference entries.
    Returns list of dicts: {sentence, citation, reference_text}
    """
    pairs = []
    for item in sentences_with_cites:
        for author, year in item['citations']:
            surname = author.split(' et al')[0].split(' & ')[0].split(',')[0].strip()
            surname = surname.split()[-1] if ' ' in surname else surname

            ref_text = ref_entries.get((surname, year))

            if not ref_text:
                base_year = re.sub(r'[a-z]$', '', year)
                if base_year != year:
                    ref_text = ref_entries.get((surname, base_year))

            if ref_text:
                pairs.append({
                    'sentence': item['sentence'],
                    'citation': f"{author} ({year})",
                    'reference_text': ref_text
                })

    return pairs


def extract_title_from_ref(ref_text):
    """Extract the title from an APA reference entry."""
    m = re.search(r'\(\d{4}[a-z]?\)\.\s*(.+?)(?:\.\s|$)', ref_text)
    if m:
        title = m.group(1).strip().rstrip('.')
        title = re.sub(r'[_*]', '', title)
        return title
    return None


def main():
    if len(sys.argv) < 2:
        print("Usage: py ref_context.py <input.docx>")
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    json_output_path = input_path.parent / f"{input_path.stem}_PAIRS.json"

    print(f"Reading: {input_path}")
    doc = Document(str(input_path))

    # Find references section
    ref_para_idx = find_references_heading_index(doc)
    print(f"References section starts at paragraph {ref_para_idx} of {len(doc.paragraphs)}")

    # Build body text
    body_text = '\n\n'.join(p.text for p in doc.paragraphs[:ref_para_idx])

    # Extract sentences with citations
    sentences_with_cites = extract_sentences_with_citations(body_text)
    total_citations = sum(len(s['citations']) for s in sentences_with_cites)
    print(f"Found {len(sentences_with_cites)} sentences with {total_citations} citations")

    # Extract reference entries
    ref_entries = extract_reference_entries(doc, ref_para_idx)
    print(f"Found {len(ref_entries)} reference entries")

    # Build matched pairs
    pairs = build_citation_ref_pairs(sentences_with_cites, ref_entries)
    print(f"Matched {len(pairs)} citation-reference pairs")

    if not pairs:
        print("\nNo matched pairs found. Run ref-check first to verify citation-reference matching.")
        sys.exit(0)

    # Deduplicate references and extract titles for search queries
    unique_refs = {}
    for p in pairs:
        ref = p['reference_text']
        if ref not in unique_refs:
            title = extract_title_from_ref(ref)
            surname = ref.split(',')[0].strip()
            year_m = re.search(r'\((\d{4})', ref)
            year = year_m.group(1) if year_m else ''
            unique_refs[ref] = {
                'title': title,
                'surname': surname,
                'year': year,
                'search_query': f"{surname} {year} {title}" if title else f"{surname} {year}"
            }

    # Build output JSON
    output = {
        'paper': input_path.name,
        'stats': {
            'sentences_with_citations': len(sentences_with_cites),
            'total_citations': total_citations,
            'reference_entries': len(ref_entries),
            'matched_pairs': len(pairs),
            'unique_references': len(unique_refs)
        },
        'pairs': pairs,
        'unique_references': {
            ref_text: info for ref_text, info in unique_refs.items()
        }
    }

    # Save JSON
    with open(json_output_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, indent=2, ensure_ascii=False)

    print(f"\nExtraction complete!")
    print(f"  Pairs: {len(pairs)}")
    print(f"  Unique references to search: {len(unique_refs)}")
    print(f"  Output: {json_output_path}")


if __name__ == "__main__":
    main()
