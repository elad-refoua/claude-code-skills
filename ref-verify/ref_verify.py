"""
ref_verify.py - Extract references from a .docx academic paper for verification + Zotero export.

Usage:
    py ref_verify.py "paper.docx"                     # Extract refs → JSON + RIS
    py ref_verify.py "paper.docx" --batch-size 10      # Custom batch size
    py ref_verify.py "paper.docx" --ris-only            # Only generate RIS (no JSON)
    py ref_verify.py --from-json "results.json"         # Generate RIS from existing JSON

Output:
    paper_REFS_TO_VERIFY.json   - Batched refs for sub-agent verification
    paper_REFS.ris              - RIS file importable by Zotero/Mendeley/EndNote
"""
import json
import re
import sys
import io
from pathlib import Path

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")


# ---------------------------------------------------------------------------
# Reference extraction
# ---------------------------------------------------------------------------

def find_references_section(doc):
    """Find the start of the References section."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if text in ("references", "reference list", "bibliography"):
            return i + 1
    return None


def extract_references(docx_path):
    """Extract all reference entries from the document."""
    doc = Document(docx_path)
    ref_start = find_references_section(doc)
    if ref_start is None:
        print("ERROR: Could not find References section!")
        sys.exit(1)

    refs = []
    for para in doc.paragraphs[ref_start:]:
        text = para.text.strip().replace("\u200f", "").replace("\u200e", "")
        if text and len(text) > 15:
            refs.append(text)
    return refs


# ---------------------------------------------------------------------------
# APA reference parser
# ---------------------------------------------------------------------------

def parse_apa_reference(text):
    """Parse an APA reference into structured fields."""
    text = text.strip()
    r = {
        "authors": "", "year": "", "title": "", "type": "unknown",
        "journal": "", "volume": "", "issue": "", "pages": "",
        "publisher": "", "editors": "", "edition": "", "doi": "",
        "full_text": text,
    }

    # Extract DOI
    doi_m = re.search(r"(https?://doi\.org/\S+)", text)
    if doi_m:
        r["doi"] = doi_m.group(1).rstrip(".")
        text_no_doi = text[:doi_m.start()].strip().rstrip(".")
    else:
        text_no_doi = text.rstrip(".")

    # Extract authors and year
    eds_m = re.match(r"^(.+?)\s*\(Eds?\.\)\.\s*\((\d{4}[a-z]?)\)\.?\s*(.*)$", text_no_doi)
    reg_m = re.match(r"^(.+?)\s*\((\d{4}[a-z]?(?:/\d{4})?)\)\.?\s*(.*)$", text_no_doi)

    if eds_m:
        r["authors"] = eds_m.group(1).strip().rstrip(",")
        r["editors"] = r["authors"]
        r["year"] = eds_m.group(2)
        remainder = eds_m.group(3).strip()
        r["type"] = "edited_book"
    elif reg_m:
        r["authors"] = reg_m.group(1).strip().rstrip(",")
        r["year"] = reg_m.group(2)
        remainder = reg_m.group(3).strip()
    else:
        return r

    if not remainder:
        return r

    # Chapter (has "In " + editors)
    in_m = re.search(r"\.\s+In\s+", remainder)
    if in_m and r["type"] != "edited_book":
        r["type"] = "chapter"
        r["title"] = remainder[:in_m.start()].strip().rstrip(".")
        after_in = remainder[in_m.end():].strip()
        eds_in = re.match(r"(.+?)\s*\(Eds?\.(?:\s*&\s*Trans\.)?\)\s*,?\s*(.*)$", after_in)
        if eds_in:
            r["editors"] = eds_in.group(1).strip()
            book_rest = eds_in.group(2).strip()
        else:
            book_rest = after_in
        paren_m = re.match(r"([^(]+?)(?:\s*\((.+?)\))?\.\s*(.*)", book_rest)
        if paren_m:
            paren_info = paren_m.group(2) or ""
            if paren_info:
                pp = re.search(r"pp?\.?\s*([\d\u2013\u2014\-]+)", paren_info)
                if pp:
                    r["pages"] = pp.group(1)
            after_book = paren_m.group(3).strip()
            if after_book:
                r["publisher"] = after_book.rstrip(".")
    elif r["type"] == "edited_book":
        title_m = re.match(r"(.+?)(?:\s*\((.+?)\))?\.\s*(.*)", remainder)
        if title_m:
            r["title"] = title_m.group(1).strip().rstrip(".")
            paren = title_m.group(2) or ""
            if "ed." in paren.lower():
                r["edition"] = paren
            r["publisher"] = title_m.group(3).strip().rstrip(".")
    else:
        # Journal or book
        j_m = re.match(
            r"(.+?)\.\s+([A-Z][^,]+?),\s*(\d+)\s*(?:\((\d+)\))?\s*,\s*([\d\u2013\u2014\-]+(?:\u2013\d+)?)\s*\.?$",
            remainder,
        )
        j_m2 = re.match(
            r"(.+?)\.\s+([A-Z][^,]+?:\s*[^,]+?),\s*(\d+)\s*(?:\((\d+)\))?\s*,\s*([\d\u2013\u2014\-]+)\s*\.?$",
            remainder,
        ) if not j_m else None

        jm = j_m or j_m2
        if jm:
            r["type"] = "journal"
            r["title"] = jm.group(1).strip()
            r["journal"] = jm.group(2).strip()
            r["volume"] = jm.group(3)
            r["issue"] = jm.group(4) or ""
            r["pages"] = jm.group(5)
        else:
            r["type"] = "book"
            b_m = re.match(r"(.+?)(?:\s*\((.+?)\))?\.\s*(.*)", remainder)
            if b_m:
                r["title"] = b_m.group(1).strip().rstrip(".")
                paren = b_m.group(2) or ""
                rest = b_m.group(3).strip().rstrip(".")
                if paren and "ed." in paren.lower():
                    r["edition"] = paren
                if rest:
                    r["publisher"] = rest.split(".")[0].strip().rstrip(".")
            else:
                r["title"] = remainder

    return r


def parse_author_year(ref_text):
    """Quick extract of first author surname and year for labeling."""
    m = re.match(r"^(.+?)\s*\((\d{4}[a-z]?)\)", ref_text)
    if not m:
        m = re.match(r"^(.+?)\s*\(Eds?\.\)\.\s*\((\d{4}[a-z]?)\)", ref_text)
    if m:
        return m.group(1).split(",")[0].strip(), m.group(2)
    return "Unknown", "????"


# ---------------------------------------------------------------------------
# RIS export (Zotero-compatible)
# ---------------------------------------------------------------------------

RIS_TYPE_MAP = {
    "journal": "JOUR",
    "book": "BOOK",
    "chapter": "CHAP",
    "edited_book": "EDBOOK",
    "unknown": "GEN",
}


def parse_authors_list(authors_str):
    """Split APA author string into individual 'Surname, Initials' entries."""
    authors_str = re.sub(r"\s*&\s*", ", ", authors_str)
    parts = re.split(r",\s*(?=[A-Z])", authors_str)
    authors = []
    current = ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r"^[A-Z]\.\s*[A-Z]?\.*$", part) or re.match(r"^[A-Z]\.$", part):
            current += ", " + part
            authors.append(current.strip())
            current = ""
        else:
            if current:
                authors.append(current.strip())
            current = part
    if current:
        authors.append(current.strip())
    return authors


def ref_to_ris(parsed, ref_num):
    """Convert a parsed reference dict to RIS format string."""
    lines = []
    ris_type = RIS_TYPE_MAP.get(parsed["type"], "GEN")
    lines.append(f"TY  - {ris_type}")
    lines.append(f"ID  - ref{ref_num:03d}")

    # Authors
    if parsed["authors"]:
        for au in parse_authors_list(parsed["authors"]):
            lines.append(f"AU  - {au}")

    # Year
    if parsed["year"]:
        lines.append(f"PY  - {parsed['year']}")

    # Title
    if parsed["title"]:
        tag = "T1" if ris_type == "JOUR" else "TI"
        lines.append(f"{tag}  - {parsed['title']}")

    # Journal
    if parsed["journal"]:
        lines.append(f"JO  - {parsed['journal']}")

    # Volume, Issue, Pages
    if parsed["volume"]:
        lines.append(f"VL  - {parsed['volume']}")
    if parsed["issue"]:
        lines.append(f"IS  - {parsed['issue']}")
    if parsed["pages"]:
        pages = parsed["pages"].replace("\u2013", "-").replace("\u2014", "-")
        if "-" in pages:
            sp, ep = pages.split("-", 1)
            lines.append(f"SP  - {sp}")
            lines.append(f"EP  - {ep}")
        else:
            lines.append(f"SP  - {pages}")

    # Publisher
    if parsed["publisher"]:
        lines.append(f"PB  - {parsed['publisher']}")

    # Editors
    if parsed["editors"] and parsed["type"] in ("chapter", "edited_book"):
        for ed in parse_authors_list(parsed["editors"]):
            lines.append(f"ED  - {ed}")

    # Edition
    if parsed["edition"]:
        lines.append(f"ET  - {parsed['edition']}")

    # DOI
    if parsed["doi"]:
        lines.append(f"DO  - {parsed['doi'].replace('https://doi.org/', '')}")
        lines.append(f"UR  - {parsed['doi']}")

    # Notes: full original text
    lines.append(f"N1  - Original APA text: {parsed['full_text']}")

    lines.append("ER  - ")
    return "\n".join(lines)


def export_ris(parsed_refs, output_path):
    """Write all parsed references as a .ris file."""
    with open(output_path, "w", encoding="utf-8") as f:
        for i, ref in enumerate(parsed_refs, 1):
            f.write(ref_to_ris(ref, i))
            f.write("\n\n")
    print(f"RIS saved: {output_path} ({len(parsed_refs)} references)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    # Handle --from-json mode
    if "--from-json" in sys.argv:
        idx = sys.argv.index("--from-json")
        if idx + 1 >= len(sys.argv):
            print("Usage: py ref_verify.py --from-json <results.json>")
            sys.exit(1)
        json_path = Path(sys.argv[idx + 1])
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)
        parsed = []
        for ref in data["all_references"]:
            parsed.append(parse_apa_reference(ref["text"]))
        ris_path = json_path.parent / f"{json_path.stem.replace('_REFS_TO_VERIFY', '')}_REFS.ris"
        export_ris(parsed, str(ris_path))
        return

    if len(sys.argv) < 2 or sys.argv[1].startswith("--"):
        print("Usage: py ref_verify.py <paper.docx> [--batch-size N] [--ris-only]")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}")
        sys.exit(1)

    batch_size = 10
    if "--batch-size" in sys.argv:
        idx = sys.argv.index("--batch-size")
        if idx + 1 < len(sys.argv):
            batch_size = int(sys.argv[idx + 1])

    ris_only = "--ris-only" in sys.argv

    print(f"Reading: {docx_path.name}")
    ref_texts = extract_references(str(docx_path))
    print(f"Found {len(ref_texts)} references")

    # Parse all references
    references = []
    parsed_refs = []
    for i, text in enumerate(ref_texts, 1):
        parsed = parse_apa_reference(text)
        parsed_refs.append(parsed)
        author, year = parse_author_year(text)
        references.append({
            "num": i,
            "label": f"{author} ({year})",
            "text": text,
        })
        print(f"  {i:2d}. [{parsed['type']:<12s}] {author} ({year})")

    # Always generate RIS
    ris_path = docx_path.parent / f"{docx_path.stem}_REFS.ris"
    export_ris(parsed_refs, str(ris_path))

    if ris_only:
        return

    # Generate JSON for sub-agent verification
    num_batches = (len(references) + batch_size - 1) // batch_size
    batches = []
    for b in range(num_batches):
        start = b * batch_size
        end = min(start + batch_size, len(references))
        batch_refs = references[start:end]
        batches.append({
            "batch_num": b + 1,
            "ref_start": batch_refs[0]["num"],
            "ref_end": batch_refs[-1]["num"],
            "count": len(batch_refs),
            "references": batch_refs,
        })

    output = {
        "paper": docx_path.stem,
        "total_references": len(references),
        "batch_size": batch_size,
        "num_batches": num_batches,
        "batches": batches,
        "all_references": references,
    }

    out_path = docx_path.parent / f"{docx_path.stem}_REFS_TO_VERIFY.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"\nJSON: {out_path.name}")
    print(f"RIS:  {ris_path.name}")
    print(f"Batches: {num_batches} (batch size {batch_size})")
    print(f"\nReady for Claude Code sub-agent verification.")


if __name__ == "__main__":
    main()
