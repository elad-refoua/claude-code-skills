"""
Reference Checker: Cross-reference in-text citations vs. reference list.
Produces a color-highlighted Word document + JSON results for LLM verification.

Usage: py ref_check.py <input.docx>
Output: <input>_REF_CHECK.docx  (highlighted Word doc)
        <input>_RESULTS.json   (structured data for Claude Code sub-agents)

Colors:
  Body citations:  GREEN = exact match,  CYAN = fuzzy match,  YELLOW = missing from refs
  References:      GREEN = cited,  CYAN = fuzzy match,  RED = not cited

This script does regex extraction + highlighting only.
Claude Code orchestrates LLM verification via sub-agents (Sonnet + Opus).
"""

import re
import sys
import copy
import json
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Highlight helpers
# ---------------------------------------------------------------------------

GREEN_HIGHLIGHT = "green"
YELLOW_HIGHLIGHT = "yellow"
RED_HIGHLIGHT = "red"
CYAN_HIGHLIGHT = "cyan"


def set_highlight(run, color_index):
    """Set highlight color on a run using Word's highlight XML."""
    rPr = run._r.get_or_add_rPr()
    highlight = rPr.find(qn("w:highlight"))
    if highlight is None:
        highlight = OxmlElement("w:highlight")
        rPr.append(highlight)
    highlight.set(qn("w:val"), color_index)


def remove_highlight(run):
    """Remove highlight from a run."""
    rPr = run._r.get_or_add_rPr()
    highlight = rPr.find(qn("w:highlight"))
    if highlight is not None:
        rPr.remove(highlight)


# ---------------------------------------------------------------------------
# Fuzzy year matching
# ---------------------------------------------------------------------------

def strip_year_suffix(year):
    """Strip letter suffix from year: '1912a' -> '1912'."""
    return re.sub(r'[a-z]$', '', year)


def build_fuzzy_lookup(item_set):
    """
    Build a fuzzy lookup from a set of (surname, year) tuples.
    Returns a dict: base_key -> list of exact keys.
    E.g., ("Freud", "1912") -> [("Freud", "1912"), ("Freud", "1912a"), ("Freud", "1912b")]
    """
    lookup = {}
    for surname, year in item_set:
        base = (surname, strip_year_suffix(year))
        if base not in lookup:
            lookup[base] = []
        lookup[base].append((surname, year))
    return lookup


def fuzzy_match_citation(surname, year, ref_set, ref_fuzzy_lookup):
    """
    Try to match a citation (surname, year) against the ref set.
    Returns: ('exact', ref_key) | ('fuzzy', ref_key) | ('none', None)
    """
    # 1. Exact match
    if (surname, year) in ref_set:
        return ('exact', (surname, year))

    # 2. Strip suffix: citation has 'a'/'b', ref doesn't
    base_year = strip_year_suffix(year)
    if base_year != year and (surname, base_year) in ref_set:
        return ('fuzzy', (surname, base_year))

    # 3. Ref has suffix, citation doesn't - check fuzzy lookup
    base_key = (surname, base_year)
    if base_key in ref_fuzzy_lookup:
        candidates = ref_fuzzy_lookup[base_key]
        # If any candidate matches (even with different suffix), it's fuzzy
        for ref_surname, ref_year in candidates:
            if ref_year != year:
                return ('fuzzy', (ref_surname, ref_year))

    return ('none', None)


def fuzzy_match_reference(surname, year, citation_set, citation_fuzzy_lookup):
    """
    Try to match a reference (surname, year) against the citation set.
    Returns: ('exact', cite_key) | ('fuzzy', cite_key) | ('none', None)
    """
    # 1. Exact match
    if (surname, year) in citation_set:
        return ('exact', (surname, year))

    # 2. Strip suffix: ref has 'a'/'b', citation doesn't
    base_year = strip_year_suffix(year)
    if base_year != year and (surname, base_year) in citation_set:
        return ('fuzzy', (surname, base_year))

    # 3. Citation has suffix, ref doesn't
    base_key = (surname, base_year)
    if base_key in citation_fuzzy_lookup:
        candidates = citation_fuzzy_lookup[base_key]
        for cite_surname, cite_year in candidates:
            if cite_year != year:
                return ('fuzzy', (cite_surname, cite_year))

    return ('none', None)


# ---------------------------------------------------------------------------
# Citation extraction from plain text
# ---------------------------------------------------------------------------

NOISE_WORDS = {
    'approaches', 'therapy', 'psychology', 'modes', 'vol', 'eds',
    'chapter', 'section', 'page', 'table', 'figure', 'note', 'press',
    'cognitive', 'relational', 'behavioral', 'emotional', 'self',
    'individual', 'interpersonal', 'focused', 'based', 'oriented',
    'and', 'the', 'for', 'with', 'from', 'into', 'between',
    'also', 'cited', 'review', 'example', 'well', 'see', 'new',
    'other', 'their', 'these', 'those', 'such', 'both', 'each',
    'psychotherapy', 'attachment', 'schema', 'technique', 'model',
    'clinical', 'therapeutic', 'treatment', 'practice', 'research',
    'humanistic', 'emotion-focused', 'self-psychology', 'self-psychology:',
    'psychoanalytic', 'existential', 'gestalt', 'integrative',
    'dynamic', 'experiential', 'systemic', 'narrative', 'dialectical',
}


def normalize_text(text):
    """Normalize text: curly quotes to straight, special chars. Removes zero-width chars."""
    # Curly/smart quotes to straight
    text = text.replace('\u2018', "'").replace('\u2019', "'")  # single quotes
    text = text.replace('\u201c', '"').replace('\u201d', '"')  # double quotes
    # Unicode hyphens/dashes to ASCII hyphen
    text = text.replace('\u2010', '-').replace('\u2011', '-')  # hyphens
    text = text.replace('\u2012', '-').replace('\u2013', '-')  # figure dash, en-dash
    # Invisible chars (REMOVES - changes length, only use for set-building)
    text = text.replace('\u200f', '').replace('\u200e', '').replace('\u200b', '')
    return text


def normalize_for_highlight(text):
    """Normalize text preserving character positions (1:1 replacements only)."""
    # Curly/smart quotes to straight (1:1, preserves positions)
    text = text.replace('\u2018', "'").replace('\u2019', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    # Unicode hyphens/dashes to ASCII hyphen (1:1)
    text = text.replace('\u2010', '-').replace('\u2011', '-')
    text = text.replace('\u2012', '-').replace('\u2013', '-')
    # DO NOT remove zero-width chars - they shift positions
    return text


def clean_author(author_str):
    """Remove prefixes and normalize an author string."""
    s = author_str.strip()
    s = re.sub(
        r'^(e\.g\.,?\s*|see\s+|cf\.\s*|for\s+review,?\s*see\s*|also\s+|as cited in\s+|as well as\s+)',
        '', s, flags=re.IGNORECASE
    ).strip()
    s = s.rstrip(',').strip()
    # Remove possessive (both straight and curly quotes)
    s = re.sub(r"['\u2019]s$", '', s)
    return s


def first_surname(author_str):
    """Extract the first author's surname from a citation author string.
    Always returns with first letter capitalized."""
    s = author_str.split(' et al')[0]
    s = s.split(' & ')[0].split(' and ')[0].split(',')[0].strip()
    s = s.strip().rstrip('.').strip()
    parts = s.split()
    if not parts:
        return normalize_surname(s)
    # If multiple capitalized words, take the LAST one (skip first names)
    # E.g., "Otto Kernberg" -> "Kernberg", "Clara Thompson" -> "Thompson"
    # BUT preserve hyphenated names: "Kabat-Zinn" stays as one token
    cap_parts = [p for p in parts if p and p[0].isupper()]
    if len(cap_parts) >= 2:
        last = cap_parts[-1]
        # Handle common "van/de/den/der" prefixes
        idx = parts.index(last) if last in parts else -1
        if idx > 0:
            prev = parts[idx - 1].lower()
            if prev in ('van', 'de', 'den', 'der', 'von', 'di', 'la', 'le', 'el', 'al'):
                return f"{parts[idx-1]} {last}"
        return normalize_surname(last)
    return normalize_surname(parts[0])


def is_noise(surname):
    """Check if a surname is a noise word (not a real author)."""
    return surname.lower() in NOISE_WORDS or len(surname) <= 2


def normalize_surname(s):
    """Normalize surname for matching: capitalize first letter."""
    s = s.strip()
    if not s:
        return s
    return s[0].upper() + s[1:] if len(s) > 1 else s.upper()


def extract_citations_from_text(plain_text):
    """
    Extract all (first_surname, year) pairs from the body text.
    Returns a set of (surname, year) tuples.
    """
    citations = set()
    text = normalize_text(plain_text)

    # 1. Parenthetical: (Author, Year) or (Author, Year; Author2, Year2)
    #    Supports author inheritance: (Gelso, 2009; 2014) → both are Gelso
    for block in re.findall(r'\(([^)]+)\)', text):
        last_author = None  # Track author for inheritance
        for part in block.split(';'):
            part = part.strip()
            years = re.findall(r'(\d{4}[a-z]?)', part)
            if years:
                author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                author_part = clean_author(author_part)
                if author_part and len(author_part) > 2:
                    surname = first_surname(author_part)
                    if not is_noise(surname):
                        last_author = surname
                        for y in years:
                            citations.add((surname, y))
                elif last_author:
                    # Bare years - inherit previous author
                    for y in years:
                        citations.add((last_author, y))

    # 2. Narrative: Author (Year) or Author et al. (Year)
    #    Handles optional first name, hyphenated names (Kabat-Zinn),
    #    and optional prefix like e.g., see, cf. before year: Bion (e.g., 1962)
    for m in re.finditer(
        r"(?:([A-Z][a-z]+(?:-[A-Z][a-z]+)?)\s+)?([A-Z][a-z]+(?:-[A-Z][a-z]+)?(?:\s+(?:&|and)\s+[A-Z][a-z]+(?:-[A-Z][a-z]+)?)?(?:\s+et\s+al\.?)?)\s*\((?:e\.g\.,?\s*|see\s+|cf\.?\s*)?(\d{4}[a-z]?(?:,\s*\d{4}[a-z]?)*)\)",
        text
    ):
        # Group 1 = optional first name, Group 2 = surname/author part, Group 3 = years
        author = m.group(2).strip()
        year_str = m.group(3)
        surname = first_surname(author)
        if not is_noise(surname):
            for y in re.findall(r'(\d{4}[a-z]?)', year_str):
                citations.add((surname, y))

    # 3. Possessive with complex parenthetical:
    #    Sullivan's (1953; as well as Erich Fromm's, 1941)
    #    Otto Kernberg's transference-focused approach (1975, 1984, 2004)
    #    Negative lookahead prevents crossing another Author's boundary
    #    (e.g., "Ferenczi's ... and Sullivan's (1953)" matches Sullivan, not Ferenczi)
    for m in re.finditer(
        r"(?:([A-Z][a-z]+)\s+)?([A-Z][a-z]+)'s\s+(?:(?![A-Z][a-z]+(?:-[A-Z][a-z]+)?'s)[^(]){0,80}\(([^)]+)\)",
        text
    ):
        # Group 1 = optional first name, Group 2 = surname, Group 3 = paren content
        possessive_surname = m.group(2)
        paren_content = m.group(3)

        for part in paren_content.split(';'):
            part = part.strip()
            years = re.findall(r'(\d{4}[a-z]?)', part)
            if not years:
                continue

            # Check if this part has its own author
            author_part = re.split(r',?\s*\d{4}', part)[0].strip()
            author_part = clean_author(author_part)

            if author_part and len(author_part) > 2:
                surname = first_surname(author_part)
                if not is_noise(surname):
                    for y in years:
                        citations.add((surname, y))
            else:
                # Bare years - inherit the possessive author
                if not is_noise(possessive_surname):
                    for y in years:
                        citations.add((possessive_surname, y))

    # 4a. Bracket citation: Author [Year], Author & Author2 [Year],
    #     or Author, Author2, & Author3 [Year] (take first surname)
    for m in re.finditer(
        r'([A-Z][a-z]+(?:-[A-Z][a-z]+)?(?:,?\s*&?\s*[A-Z][a-z]+(?:-[A-Z][a-z]+)?)*(?:\s+et\s+al\.?)?)\s*\[(\d{4}[a-z]?)\]',
        text
    ):
        author = m.group(1).strip()
        year = m.group(2)
        surname = first_surname(clean_author(author))
        if not is_noise(surname):
            citations.add((surname, year))

    # 4b. Multi-citation bracket block: [Author1, Year1; Author2, Year2]
    for block_match in re.finditer(r'\[([^\]]*\d{4}[^\]]*;[^\]]*)\]', text):
        block = block_match.group(1)
        for part in block.split(';'):
            part = part.strip()
            years = re.findall(r'(\d{4}[a-z]?)', part)
            if years:
                author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                author_part = clean_author(author_part)
                if author_part and len(author_part) > 2:
                    surname = first_surname(author_part)
                    if not is_noise(surname):
                        for y in years:
                            citations.add((surname, y))

    # 5. Narrative with "and colleagues/coworkers/collaborators":
    #    "Stiles and colleagues (e.g., 2009; Kramer & Stiles, 2015)"
    for m in re.finditer(
        r'([A-Z][a-z]+)\s+and\s+(?:colleagues|coworkers|collaborators)\s*\(([^)]+)\)',
        text
    ):
        narrative_surname = m.group(1)
        paren_content = m.group(2)
        for part in paren_content.split(';'):
            part = part.strip()
            years = re.findall(r'(\d{4}[a-z]?)', part)
            if not years:
                continue
            author_part = re.split(r',?\s*\d{4}', part)[0].strip()
            author_part = clean_author(author_part)
            if author_part and len(author_part) > 2:
                surname = first_surname(author_part)
                if not is_noise(surname):
                    for y in years:
                        citations.add((surname, y))
            else:
                if not is_noise(narrative_surname):
                    for y in years:
                        citations.add((narrative_surname, y))

    # 6. Narrative with "et al." followed by complex parenthetical:
    #    "Author et al. (e.g., 2009; Other, 2015)"
    for m in re.finditer(
        r'([A-Z][a-z]+)\s+et\s+al\.?\s*\(([^)]+)\)',
        text
    ):
        narrative_surname = m.group(1)
        paren_content = m.group(2)
        for part in paren_content.split(';'):
            part = part.strip()
            years = re.findall(r'(\d{4}[a-z]?)', part)
            if not years:
                continue
            author_part = re.split(r',?\s*\d{4}', part)[0].strip()
            author_part = clean_author(author_part)
            if author_part and len(author_part) > 2:
                surname = first_surname(author_part)
                if not is_noise(surname):
                    for y in years:
                        citations.add((surname, y))
            else:
                if not is_noise(narrative_surname):
                    for y in years:
                        citations.add((narrative_surname, y))

    # 7. Catch lowercase citations: "stiles (2009)" - normalize to "Stiles"
    for m in re.finditer(
        r'([a-z][a-z]+(?:-[a-zA-Z]+)?)\s*[\(,]\s*(\d{4}[a-z]?)',
        text
    ):
        surname = normalize_surname(m.group(1))
        year = m.group(2)
        if not is_noise(surname):
            citations.add((surname, year))

    # Normalize all surnames in the set (capitalize first letter)
    normalized = set()
    for surname, year in citations:
        normalized.add((normalize_surname(surname), year))

    return normalized


def extract_references_from_doc(doc, ref_para_idx):
    """
    Extract (first_surname, year) pairs from reference paragraphs in the doc.
    Returns a set of (surname, year) tuples.
    Also auto-assigns a/b suffixes when duplicate (surname, year) pairs exist.
    """
    raw_refs = []

    for pi in range(ref_para_idx + 1, len(doc.paragraphs)):
        text = normalize_text(doc.paragraphs[pi].text.strip())
        if not text:
            continue
        if text in ('References', 'REFERENCES', 'Reference List', 'Bibliography'):
            continue

        year_match = re.search(r'\((\d{4}[a-z]?)\)', text)
        if year_match:
            year = year_match.group(1)
            before = text[:text.index('(')].strip().rstrip(',').strip()
            parts = before.split(',')[0].strip().rstrip('.').split()
            surname = parts[0] if parts else ''
            if surname and len(surname) > 1:
                raw_refs.append((surname, year, pi))

    # Auto-assign a/b suffixes for duplicates
    refs = set()
    # Count occurrences of each (surname, base_year)
    from collections import Counter
    base_counts = Counter()
    for surname, year, pi in raw_refs:
        base = (surname, strip_year_suffix(year))
        base_counts[base] += 1

    # Track suffix assignment
    suffix_tracker = {}
    for surname, year, pi in raw_refs:
        base = (surname, strip_year_suffix(year))
        if base_counts[base] > 1 and not re.search(r'[a-z]$', year):
            # Duplicate base year without suffix - assign one
            if base not in suffix_tracker:
                suffix_tracker[base] = ord('a')
            suffix = chr(suffix_tracker[base])
            suffix_tracker[base] += 1
            refs.add((surname, year + suffix))
            # Also keep the base year version for fuzzy matching
            refs.add((surname, year))
        else:
            refs.add((surname, year))

    return refs


# ---------------------------------------------------------------------------
# Document-level processing: find & highlight citation runs
# ---------------------------------------------------------------------------

def get_body_paragraphs(doc, ref_para_idx):
    """
    All paragraphs before References heading, including inside tables.
    Uses XML element identity as stop sentinel to avoid index mismatch.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    stop_elem = doc.paragraphs[ref_para_idx]._element if ref_para_idx < len(doc.paragraphs) else None

    result = []
    for child in doc.element.body.iterchildren():
        if child is stop_elem:
            break
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            result.append(Paragraph(child, doc))
        elif tag == 'tbl':
            tbl = Table(child, doc)
            seen = set()
            for row in tbl.rows:
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid not in seen:
                        seen.add(cid)
                        result.extend(cell.paragraphs)
    return result


def get_plain_text(doc):
    """Get all text from document paragraphs."""
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    return '\n'.join(parts)


def find_references_heading_index(doc):
    """Find the paragraph index where 'References' heading starts."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in ('References', 'REFERENCES', 'Reference List', 'Bibliography'):
            return i
    return len(doc.paragraphs)


def determine_highlight_color(surname, year, ref_set, ref_fuzzy_lookup):
    """Determine the highlight color for a citation based on match type."""
    match_type, _ = fuzzy_match_citation(surname, year, ref_set, ref_fuzzy_lookup)
    if match_type == 'exact':
        return GREEN_HIGHLIGHT
    elif match_type == 'fuzzy':
        return CYAN_HIGHLIGHT
    else:
        return YELLOW_HIGHLIGHT


def highlight_body_citations(doc, ref_para_idx, citation_set, ref_set, ref_fuzzy_lookup,
                             add_comments=False):
    """
    Walk through body paragraphs (incl. tables), find citation spans, highlight them.
    Optionally adds Word bubble comments for non-green items.
    Returns counts and lists.
    """
    matched = 0
    fuzzy_matched = 0
    unmatched = 0
    unmatched_list = []
    fuzzy_list = []

    for para in get_body_paragraphs(doc, ref_para_idx):
        full_text = normalize_for_highlight(para.text)
        if not full_text.strip():
            continue

        # Track already-highlighted spans to avoid double-processing
        highlighted_spans = []

        # --- 4a. Bracket citation: Author [Year] or Author, Author2, & Author3 [Year] ---
        for m in re.finditer(
            r'([A-Z][a-z]+(?:-[A-Z][a-z]+)?(?:,?\s*&?\s*[A-Z][a-z]+(?:-[A-Z][a-z]+)?)*(?:\s+et\s+al\.?)?)\s*\[(\d{4}[a-z]?)\]',
            full_text
        ):
            span_start = m.start()
            span_end = m.end()
            author = m.group(1).strip()
            year = m.group(2)
            surname = first_surname(clean_author(author))

            if is_noise(surname):
                continue

            mt, _ = fuzzy_match_citation(surname, year, ref_set, ref_fuzzy_lookup)
            if mt == 'exact':
                color = GREEN_HIGHLIGHT
                matched += 1
            elif mt == 'fuzzy':
                color = CYAN_HIGHLIGHT
                fuzzy_matched += 1
                fuzzy_list.append(f"{surname} ({year})")
            else:
                color = YELLOW_HIGHLIGHT
                unmatched += 1
                unmatched_list.append(f"{surname} ({year})")

            highlight_span_in_paragraph(para, span_start, span_end, color)
            highlighted_spans.append((span_start, span_end))
            if add_comments and color == YELLOW_HIGHLIGHT:
                add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

        # --- 4b. Multi-citation bracket block: [Author, Year; Author2, Year2] ---
        for m in re.finditer(r'\[([^\]]*\d{4}[^\]]*;[^\]]*)\]', full_text):
            span_start = m.start()
            span_end = m.end()

            # Skip if overlaps with already-highlighted span (e.g., 4a already handled)
            if any(hs <= span_start < he or hs < span_end <= he for hs, he in highlighted_spans):
                continue

            block = m.group(1)
            block_citations = []

            for part in block.split(';'):
                part = part.strip()
                years = re.findall(r'(\d{4}[a-z]?)', part)
                if years:
                    author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                    author_part = clean_author(author_part)
                    if author_part and len(author_part) > 2:
                        surname = first_surname(author_part)
                        if not is_noise(surname):
                            for y in years:
                                block_citations.append((surname, y))

            if block_citations:
                # Determine color based on worst-case match
                colors = [determine_highlight_color(s, y, ref_set, ref_fuzzy_lookup) for s, y in block_citations]
                if YELLOW_HIGHLIGHT in colors:
                    color = YELLOW_HIGHLIGHT
                elif CYAN_HIGHLIGHT in colors:
                    color = CYAN_HIGHLIGHT
                else:
                    color = GREEN_HIGHLIGHT

                for s, y in block_citations:
                    mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
                    if mt == 'exact':
                        matched += 1
                    elif mt == 'fuzzy':
                        fuzzy_matched += 1
                        fuzzy_list.append(f"{s} ({y})")
                    else:
                        unmatched += 1
                        unmatched_list.append(f"{s} ({y})")

                highlight_span_in_paragraph(para, span_start, span_end, color)
                highlighted_spans.append((span_start, span_end))
                if add_comments and color == YELLOW_HIGHLIGHT:
                    add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

        # --- Parenthetical citations: (Author, Year; Author2, Year2) ---
        #     Supports author inheritance: (Gelso, 2009; 2014) → both are Gelso
        for m in re.finditer(r'\(([^)]*\d{4}[a-z]?[^)]*)\)', full_text):
            span_start = m.start()
            span_end = m.end()

            # Skip if overlaps with already-highlighted span
            if any(hs <= span_start < he or hs < span_end <= he for hs, he in highlighted_spans):
                continue

            cite_block = m.group(1)
            block_citations = []
            last_author = None

            for part in cite_block.split(';'):
                part = part.strip()
                years = re.findall(r'(\d{4}[a-z]?)', part)
                if years:
                    author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                    author_part = clean_author(author_part)
                    if author_part and len(author_part) > 2:
                        surname = first_surname(author_part)
                        if not is_noise(surname):
                            last_author = surname
                            for y in years:
                                block_citations.append((surname, y))
                    elif last_author:
                        # Bare years - inherit previous author
                        for y in years:
                            block_citations.append((last_author, y))

            if block_citations:
                colors = [determine_highlight_color(s, y, ref_set, ref_fuzzy_lookup) for s, y in block_citations]
                if YELLOW_HIGHLIGHT in colors:
                    color = YELLOW_HIGHLIGHT
                elif CYAN_HIGHLIGHT in colors:
                    color = CYAN_HIGHLIGHT
                else:
                    color = GREEN_HIGHLIGHT

                for s, y in block_citations:
                    mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
                    if mt == 'exact':
                        matched += 1
                    elif mt == 'fuzzy':
                        fuzzy_matched += 1
                        fuzzy_list.append(f"{s} ({y})")
                    else:
                        unmatched += 1
                        unmatched_list.append(f"{s} ({y})")

                highlight_span_in_paragraph(para, span_start, span_end, color)
                highlighted_spans.append((span_start, span_end))
                if add_comments and color == YELLOW_HIGHLIGHT:
                    add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

        # --- Narrative citations: Author (Year) or Author et al. (Year) ---
        # Handles hyphenated names like Kabat-Zinn, optional first name,
        # and optional prefix like e.g., see, cf. before year: Bion (e.g., 1962)
        for m in re.finditer(
            r"(?:([A-Z][a-z]+(?:-[A-Z][a-z]+)?)\s+)?([A-Z][a-z]+(?:-[A-Z][a-z]+)?(?:\s+(?:&|and)\s+[A-Z][a-z]+(?:-[A-Z][a-z]+)?)?(?:\s+et\s+al\.?)?)\s*\((?:e\.g\.,?\s*|see\s+|cf\.?\s*)?(\d{4}[a-z]?(?:,\s*\d{4}[a-z]?)*)\)",
            full_text
        ):
            span_start = m.start()
            span_end = m.end()

            # Skip if overlaps with already-highlighted span
            if any(hs <= span_start < he or hs < span_end <= he for hs, he in highlighted_spans):
                continue

            author = m.group(2).strip()
            year_str = m.group(3)
            surname = first_surname(author)
            years_found = re.findall(r'(\d{4}[a-z]?)', year_str)

            if is_noise(surname):
                continue

            block_citations = [(surname, y) for y in years_found]
            colors = [determine_highlight_color(s, y, ref_set, ref_fuzzy_lookup) for s, y in block_citations]
            if YELLOW_HIGHLIGHT in colors:
                color = YELLOW_HIGHLIGHT
            elif CYAN_HIGHLIGHT in colors:
                color = CYAN_HIGHLIGHT
            else:
                color = GREEN_HIGHLIGHT

            for s, y in block_citations:
                mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
                if mt == 'exact':
                    matched += 1
                elif mt == 'fuzzy':
                    fuzzy_matched += 1
                    fuzzy_list.append(f"{s} ({y})")
                else:
                    unmatched += 1
                    unmatched_list.append(f"{s} ({y})")

            highlight_span_in_paragraph(para, span_start, span_end, color)
            highlighted_spans.append((span_start, span_end))
            if add_comments and color == YELLOW_HIGHLIGHT:
                add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

        # --- Possessive citations: Author's ... (Year1, Year2; OtherAuthor, Year) ---
        # Negative lookahead prevents crossing another Author's boundary
        for m in re.finditer(
            r"(?:([A-Z][a-z]+)\s+)?([A-Z][a-z]+)'s\s+(?:(?![A-Z][a-z]+(?:-[A-Z][a-z]+)?'s)[^(]){0,80}\(([^)]+)\)",
            full_text
        ):
            span_start = m.start()
            span_end = m.end()

            # Skip if overlaps
            if any(hs <= span_start < he or hs < span_end <= he for hs, he in highlighted_spans):
                continue

            possessive_surname = m.group(2)
            paren_content = m.group(3)

            # Only process if paren content contains years
            if not re.search(r'\d{4}', paren_content):
                continue

            block_citations = []
            for part in paren_content.split(';'):
                part = part.strip()
                years = re.findall(r'(\d{4}[a-z]?)', part)
                if not years:
                    continue
                author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                author_part = clean_author(author_part)
                if author_part and len(author_part) > 2:
                    surname = first_surname(author_part)
                    if not is_noise(surname):
                        for y in years:
                            block_citations.append((surname, y))
                else:
                    if not is_noise(possessive_surname):
                        for y in years:
                            block_citations.append((possessive_surname, y))

            if block_citations:
                colors = [determine_highlight_color(s, y, ref_set, ref_fuzzy_lookup) for s, y in block_citations]
                if YELLOW_HIGHLIGHT in colors:
                    color = YELLOW_HIGHLIGHT
                elif CYAN_HIGHLIGHT in colors:
                    color = CYAN_HIGHLIGHT
                else:
                    color = GREEN_HIGHLIGHT

                for s, y in block_citations:
                    mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
                    if mt == 'exact':
                        matched += 1
                    elif mt == 'fuzzy':
                        fuzzy_matched += 1
                        fuzzy_list.append(f"{s} ({y})")
                    else:
                        unmatched += 1
                        unmatched_list.append(f"{s} ({y})")

                highlight_span_in_paragraph(para, span_start, span_end, color)
                highlighted_spans.append((span_start, span_end))
                if add_comments and color == YELLOW_HIGHLIGHT:
                    add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

        # --- "Author and colleagues" pattern ---
        for m in re.finditer(
            r'([A-Z][a-z]+)\s+and\s+(?:colleagues|coworkers|collaborators)\s*\(([^)]+)\)',
            full_text
        ):
            span_start = m.start()
            span_end = m.end()

            if any(hs <= span_start < he or hs < span_end <= he for hs, he in highlighted_spans):
                continue

            narrative_surname = m.group(1)
            paren_content = m.group(2)
            block_citations = []

            for part in paren_content.split(';'):
                part = part.strip()
                years = re.findall(r'(\d{4}[a-z]?)', part)
                if not years:
                    continue
                author_part = re.split(r',?\s*\d{4}', part)[0].strip()
                author_part = clean_author(author_part)
                if author_part and len(author_part) > 2:
                    surname = first_surname(author_part)
                    if not is_noise(surname):
                        for y in years:
                            block_citations.append((surname, y))
                else:
                    if not is_noise(narrative_surname):
                        for y in years:
                            block_citations.append((narrative_surname, y))

            if block_citations:
                colors = [determine_highlight_color(s, y, ref_set, ref_fuzzy_lookup) for s, y in block_citations]
                if YELLOW_HIGHLIGHT in colors:
                    color = YELLOW_HIGHLIGHT
                elif CYAN_HIGHLIGHT in colors:
                    color = CYAN_HIGHLIGHT
                else:
                    color = GREEN_HIGHLIGHT

                for s, y in block_citations:
                    mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
                    if mt == 'exact':
                        matched += 1
                    elif mt == 'fuzzy':
                        fuzzy_matched += 1
                        fuzzy_list.append(f"{s} ({y})")
                    else:
                        unmatched += 1
                        unmatched_list.append(f"{s} ({y})")

                highlight_span_in_paragraph(para, span_start, span_end, color)
                highlighted_spans.append((span_start, span_end))
                if add_comments and color == YELLOW_HIGHLIGHT:
                    add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

        # --- Lowercase narrative citations: stiles (2009) ---
        for m in re.finditer(
            r'([a-z][a-z]+(?:-[a-zA-Z]+)?)\s*\((\d{4}[a-z]?(?:,\s*\d{4}[a-z]?)*)\)',
            full_text
        ):
            span_start = m.start()
            span_end = m.end()

            if any(hs <= span_start < he or hs < span_end <= he for hs, he in highlighted_spans):
                continue

            surname = normalize_surname(m.group(1))
            year_str = m.group(2)

            if is_noise(surname):
                continue

            years_found = re.findall(r'(\d{4}[a-z]?)', year_str)
            block_citations = [(surname, y) for y in years_found]

            colors = [determine_highlight_color(s, y, ref_set, ref_fuzzy_lookup) for s, y in block_citations]
            if YELLOW_HIGHLIGHT in colors:
                color = YELLOW_HIGHLIGHT
            elif CYAN_HIGHLIGHT in colors:
                color = CYAN_HIGHLIGHT
            else:
                color = GREEN_HIGHLIGHT

            for s, y in block_citations:
                mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
                if mt == 'exact':
                    matched += 1
                elif mt == 'fuzzy':
                    fuzzy_matched += 1
                    fuzzy_list.append(f"{s} ({y})")
                else:
                    unmatched += 1
                    unmatched_list.append(f"{s} ({y})")

            highlight_span_in_paragraph(para, span_start, span_end, color)
            highlighted_spans.append((span_start, span_end))
            if add_comments and color == YELLOW_HIGHLIGHT:
                add_comment_to_span(doc, para, span_start, span_end, "Citation not found in reference list")

    return matched, fuzzy_matched, unmatched, unmatched_list, fuzzy_list


def split_run(run, split_at):
    """
    Split a run at character position split_at.
    Returns (left_run, right_run). The original run becomes left_run.
    """
    r_elem = run._r
    right_r = copy.deepcopy(r_elem)
    full_text = run.text

    # Clear existing w:t nodes in both
    for t_elem in r_elem.findall(qn('w:t')):
        r_elem.remove(t_elem)
    for t_elem in right_r.findall(qn('w:t')):
        right_r.remove(t_elem)

    # Set left text
    left_text = full_text[:split_at]
    new_t_left = OxmlElement('w:t')
    new_t_left.text = left_text
    if left_text.startswith(' ') or left_text.endswith(' '):
        new_t_left.set(qn('xml:space'), 'preserve')
    r_elem.append(new_t_left)

    # Set right text
    right_text = full_text[split_at:]
    new_t_right = OxmlElement('w:t')
    new_t_right.text = right_text
    if right_text.startswith(' ') or right_text.endswith(' '):
        new_t_right.set(qn('xml:space'), 'preserve')
    right_r.append(new_t_right)

    # Insert right after left
    r_elem.addnext(right_r)
    from docx.text.run import Run
    return run, Run(right_r, run._element.getparent())


def highlight_span_in_paragraph(para, char_start, char_end, color_index):
    """
    Highlight characters from char_start to char_end in a paragraph.
    Splits runs at boundaries so only the exact span is highlighted.
    """
    for boundary in [char_start, char_end]:
        pos = 0
        for run in list(para.runs):
            run_len = len(run.text)
            run_start = pos
            run_end = pos + run_len
            if run_start < boundary < run_end:
                split_run(run, boundary - run_start)
                break
            pos += run_len

    pos = 0
    for run in para.runs:
        run_len = len(run.text)
        run_start = pos
        run_end = pos + run_len
        if run_start >= char_start and run_end <= char_end and run_len > 0:
            set_highlight(run, color_index)
        pos += run_len


def add_comment_to_span(doc, para, char_start, char_end, comment_text, author="ref-check"):
    """Add a Word bubble comment anchored to runs in [char_start, char_end).
    Requires python-docx >= 1.2.0. Silently skips if comment fails."""
    span_runs = []
    pos = 0
    for run in para.runs:
        rlen = len(run.text)
        if pos >= char_start and pos + rlen <= char_end and rlen > 0:
            span_runs.append(run)
        pos += rlen
    if span_runs:
        try:
            doc.add_comment(runs=span_runs, text=comment_text, author=author)
        except Exception:
            pass


def highlight_references(doc, ref_para_idx, citation_set, citation_fuzzy_lookup,
                         add_comments=False):
    """
    Highlight reference paragraphs: green if cited, cyan if fuzzy, red if not.
    Only highlights the "Author, X. (Year)." portion, not the full line.
    Optionally adds Word bubble comments for non-green items.
    """
    cited_count = 0
    fuzzy_count = 0
    uncited_count = 0
    uncited_list = []
    fuzzy_list = []

    for pi in range(ref_para_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[pi]
        text = para.text.strip()
        if not text:
            continue

        text_clean = normalize_text(text)
        year_match = re.search(r'\((\d{4}[a-z]?)\)', text_clean)
        if not year_match:
            continue

        year = year_match.group(1)
        before = text_clean[:text_clean.index('(')].strip().rstrip(',').strip()
        parts = before.split(',')[0].strip().rstrip('.').split()
        surname = parts[0] if parts else ''

        if not surname:
            continue

        match_type, _ = fuzzy_match_reference(surname, year, citation_set, citation_fuzzy_lookup)

        if match_type == 'exact':
            cited_count += 1
            color = GREEN_HIGHLIGHT
        elif match_type == 'fuzzy':
            fuzzy_count += 1
            fuzzy_list.append(f"{surname} ({year})")
            color = CYAN_HIGHLIGHT
        else:
            uncited_count += 1
            uncited_list.append(f"{surname} ({year})")
            color = RED_HIGHLIGHT

        # Only highlight from start up to and including "(Year)"
        year_paren_match = re.search(r'\(\d{4}[a-z]?\)', para.text)
        if year_paren_match:
            span_end = year_paren_match.end()
        else:
            span_end = len(para.text)

        highlight_span_in_paragraph(para, 0, span_end, color)
        if add_comments and color == RED_HIGHLIGHT:
            add_comment_to_span(doc, para, 0, span_end, "Reference not cited in body text")

    return cited_count, fuzzy_count, uncited_count, uncited_list, fuzzy_list


# ---------------------------------------------------------------------------
# Self-Learning Mechanism (file-based, no API)
# ---------------------------------------------------------------------------

LEARNED_PATTERNS_FILE = Path(__file__).parent / "learned_patterns.json"


def load_learned_patterns():
    """Load previously learned patterns from JSON file."""
    if LEARNED_PATTERNS_FILE.exists():
        try:
            with open(LEARNED_PATTERNS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Ensure required keys exist
                data.setdefault('version', 1)
                data.setdefault('noise_words', [])
                data.setdefault('cross_matches', [])
                return data
        except (json.JSONDecodeError, IOError):
            pass
    return {"version": 1, "noise_words": [], "cross_matches": []}


def save_learned_patterns(patterns):
    """Save learned patterns to JSON file."""
    with open(LEARNED_PATTERNS_FILE, 'w', encoding='utf-8') as f:
        json.dump(patterns, f, indent=2, ensure_ascii=False)


def apply_learned_noise_words(learned):
    """Add learned noise words to the global NOISE_WORDS set."""
    global NOISE_WORDS
    for word in learned.get('noise_words', []):
        NOISE_WORDS.add(word.lower())


def apply_learned_cross_matches(learned, unmatched_citations, uncited_refs):
    """
    Check if any current unmatched citation + uncited reference pair
    was previously learned as a cross-match (from a prior Opus run).
    Returns list of auto-matches: [(citation_str, ref_str, reason), ...]
    Only fires when BOTH the citation and reference exist in current paper.
    """
    auto_matches = []
    if not learned.get('cross_matches'):
        return auto_matches

    # Parse unmatched citations into (author, year) for lookup
    cite_lookup = {}
    for c in unmatched_citations:
        m = re.match(r'(\w[\w-]*)\s*\((\d{4}[a-z]?)\)', c)
        if m:
            cite_lookup[(m.group(1).lower(), m.group(2))] = c

    # Parse uncited refs into (author, year) for lookup
    ref_lookup = {}
    for r in uncited_refs:
        m = re.match(r'(\w[\w-]*)\s*\((\d{4}[a-z]?)\)', r)
        if m:
            ref_lookup[(m.group(1).lower(), m.group(2))] = r

    for cm in learned['cross_matches']:
        author = cm.get('author', '').lower()
        cite_year = cm.get('cite_year', '')
        ref_year = cm.get('ref_year', '')
        reason = cm.get('reason', 'learned from previous Opus verification')

        cite_key = (author, cite_year)
        ref_key = (author, ref_year)

        if cite_key in cite_lookup and ref_key in ref_lookup:
            auto_matches.append((
                cite_lookup[cite_key],
                ref_lookup[ref_key],
                reason
            ))

    return auto_matches


# ---------------------------------------------------------------------------
# Legend insertion
# ---------------------------------------------------------------------------

def insert_legend(doc):
    """Insert a color legend paragraph at the top of the document."""
    # Create legend text
    legend_para = doc.paragraphs[0]._element
    new_para = OxmlElement('w:p')

    # Add legend runs with colored highlights
    legend_items = [
        ("Reference Check: ", None),
        (" EXACT MATCH ", GREEN_HIGHLIGHT),
        ("  ", None),
        (" FUZZY MATCH ", CYAN_HIGHLIGHT),
        ("  ", None),
        (" MISSING FROM REFS ", YELLOW_HIGHLIGHT),
        ("  ", None),
        (" NOT CITED IN TEXT ", RED_HIGHLIGHT),
    ]

    for text, color in legend_items:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Small font
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '16')  # 8pt
        rPr.append(sz)

        if color:
            hl = OxmlElement('w:highlight')
            hl.set(qn('w:val'), color)
            rPr.append(hl)
            # Bold for legend items
            b = OxmlElement('w:b')
            rPr.append(b)

        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        new_para.append(r)

    # Insert before first paragraph
    doc.element.body.insert(0, new_para)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def add_findings_comments(docx_path, findings_json_path):
    """
    Post-processing: add Word bubble comments from sub-agent findings
    to an already-highlighted REF_CHECK docx.

    Called by Claude Code after Opus verification:
      py ref_check.py --add-comments <REF_CHECK.docx> <findings.json>

    findings.json format:
    {
      "cross_matches": [{"citation": "...", "reference": "...", "reason": "..."}],
      "false_positives": [{"text": "...", "reason": "..."}],
      "possibly_cited_refs": [{"reference": "...", "evidence": "..."}],
      "fuzzy_comments": [{"citation": "...", "comment": "..."}],
      "other_issues": ["issue1"],
      "additional": [{"text": "...", "note": "..."}]
    }
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    findings_path = Path(findings_json_path)
    if not findings_path.exists():
        print(f"[COMMENTS] Findings file not found: {findings_path}")
        return

    with open(findings_path, 'r', encoding='utf-8') as f:
        findings = json.load(f)

    doc = Document(str(docx_path))

    # Collect all paragraphs (body + table cells) for text search
    all_paragraphs = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            all_paragraphs.append(Paragraph(child, doc))
        elif tag == 'tbl':
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

    def find_and_comment(search_text, comment_text):
        """Find first paragraph containing search_text, add comment to that run."""
        for para in all_paragraphs:
            idx = para.text.find(search_text)
            if idx != -1:
                pos = 0
                for run in para.runs:
                    rlen = len(run.text)
                    if pos <= idx < pos + rlen:
                        try:
                            doc.add_comment(runs=[run], text=comment_text,
                                            author='ref-check (Opus)')
                        except Exception:
                            pass
                        return True
                    pos += rlen
        return False

    added = 0
    for cm in findings.get('cross_matches', []):
        cite = cm.get('citation', '')
        ref = cm.get('reference', '')
        reason = cm.get('reason', '')
        search = cite.split('(')[0].strip()
        if find_and_comment(search, f"Possible match: {cite} \u2194 {ref} ({reason})"):
            added += 1

    for fp in findings.get('false_positives', []):
        text = fp.get('text', '')
        reason = fp.get('reason', 'Not a real citation')
        search = text.split('(')[0].strip()
        if find_and_comment(search, f"Not a real citation: {reason}"):
            added += 1

    for fc in findings.get('fuzzy_comments', []):
        cite = fc.get('citation', '')
        comment = fc.get('comment', '')
        if cite and comment:
            search = cite.split('(')[0].strip()
            if find_and_comment(search, comment):
                added += 1

    for pr in findings.get('possibly_cited_refs', []):
        ref = pr.get('reference', '')
        evidence = pr.get('evidence', '')
        if ref and evidence:
            search = ref.split('(')[0].strip()
            if find_and_comment(search, f"Possibly cited: {evidence}"):
                added += 1

    for item in findings.get('additional', []):
        text = item.get('text', '')
        note = item.get('note', '')
        if text and note:
            search = text.split('(')[0].strip()
            if find_and_comment(search, note):
                added += 1

    doc.save(str(docx_path))
    print(f"[COMMENTS] Added {added} Opus findings comments to: {docx_path}")


def main():
    # Handle --add-comments sub-command
    if '--add-comments' in sys.argv:
        idx = sys.argv.index('--add-comments')
        if idx + 2 < len(sys.argv):
            add_findings_comments(sys.argv[idx + 1], sys.argv[idx + 2])
        else:
            print("Usage: py ref_check.py --add-comments <REF_CHECK.docx> <findings.json>")
        sys.exit(0)

    if len(sys.argv) < 2:
        print("Usage: py ref_check.py <input.docx> [--comments]")
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    output_path = input_path.parent / f"{input_path.stem}_REF_CHECK{input_path.suffix}"
    json_path = input_path.parent / f"{input_path.stem}_RESULTS.json"

    # Load learned patterns (noise words, cross-matches from prior runs)
    learned = load_learned_patterns()
    learned_noise_count = len(learned.get('noise_words', []))
    learned_cm_count = len(learned.get('cross_matches', []))
    if learned_noise_count or learned_cm_count:
        print(f"[LEARN] Loaded {learned_noise_count} noise words, {learned_cm_count} cross-matches from memory")
    apply_learned_noise_words(learned)

    # Optional flags
    add_comments = '--comments' in sys.argv
    if add_comments:
        print("[COMMENTS] Word bubble comments enabled")

    print(f"Reading: {input_path}")
    doc = Document(str(input_path))

    # Find references section boundary
    ref_para_idx = find_references_heading_index(doc)
    print(f"References section starts at paragraph {ref_para_idx} of {len(doc.paragraphs)}")

    # Build body text and reference text (before/after references heading)
    # Includes text from tables (not just doc.paragraphs)
    body_paragraphs = get_body_paragraphs(doc, ref_para_idx)
    body_text = '\n\n'.join(p.text for p in body_paragraphs)
    ref_text = '\n'.join(p.text.strip() for p in doc.paragraphs[ref_para_idx+1:] if p.text.strip())

    # Extract citation keys and reference keys
    citation_set = extract_citations_from_text(body_text)
    ref_set = extract_references_from_doc(doc, ref_para_idx)

    print(f"Found {len(citation_set)} unique in-text citations (regex)")
    print(f"Found {len(ref_set)} unique references")

    # Build fuzzy lookup tables
    ref_fuzzy_lookup = build_fuzzy_lookup(ref_set)
    citation_fuzzy_lookup = build_fuzzy_lookup(citation_set)

    # Highlight body citations
    body_matched, body_fuzzy, body_unmatched, unmatched_list, fuzzy_cite_list = \
        highlight_body_citations(doc, ref_para_idx, citation_set, ref_set, ref_fuzzy_lookup,
                                 add_comments=add_comments)

    # Highlight references
    ref_cited, ref_fuzzy, ref_uncited, uncited_list, fuzzy_ref_list = \
        highlight_references(doc, ref_para_idx, citation_set, citation_fuzzy_lookup,
                             add_comments=add_comments)

    # Apply learned cross-matches (from prior Opus runs)
    auto_matches = apply_learned_cross_matches(learned, unmatched_list, uncited_list)
    if auto_matches:
        print(f"\n[LEARN] Applied {len(auto_matches)} learned cross-match(es):")
        for cite_str, ref_str, reason in auto_matches:
            print(f"         {cite_str} <-> {ref_str} ({reason})")
            # Move from yellow/red to fuzzy (cyan) in counts
            if cite_str in unmatched_list:
                unmatched_list.remove(cite_str)
                fuzzy_cite_list.append(cite_str)
                body_unmatched -= 1
                body_fuzzy += 1
            if ref_str in uncited_list:
                uncited_list.remove(ref_str)
                fuzzy_ref_list.append(ref_str)
                ref_uncited -= 1
                ref_fuzzy += 1

    # Insert legend at top
    insert_legend(doc)

    # Build matched citations list for JSON output
    matched_citations = []
    for s, y in citation_set:
        mt, _ = fuzzy_match_citation(s, y, ref_set, ref_fuzzy_lookup)
        if mt == 'exact':
            matched_citations.append(f"{s} ({y})")

    # Save highlighted Word doc
    doc.save(str(output_path))
    print(f"\nSaved: {output_path}")

    # Save JSON results for Claude Code sub-agents
    results_json = {
        "paper": input_path.stem,
        "body_text": body_text,
        "ref_text": ref_text,
        "stats": {
            "regex_citations": len(citation_set),
            "references": len(ref_set),
            "body_matched": body_matched,
            "body_fuzzy": body_fuzzy,
            "body_unmatched": body_unmatched,
            "ref_cited": ref_cited,
            "ref_fuzzy": ref_fuzzy,
            "ref_uncited": ref_uncited,
        },
        "matched_citations": sorted(set(matched_citations)),
        "fuzzy_matches": sorted(set(fuzzy_cite_list + fuzzy_ref_list)),
        "unmatched_citations": sorted(set(unmatched_list)),
        "uncited_references": sorted(set(uncited_list)),
        "citation_set": [f"{s} ({y})" for s, y in sorted(citation_set)],
        "ref_set": [f"{s} ({y})" for s, y in sorted(ref_set)],
    }

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(results_json, f, indent=2, ensure_ascii=False)
    print(f"Saved: {json_path}")

    # Print summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"\nBody citations - exact match (GREEN):   {body_matched}")
    print(f"Body citations - fuzzy match (CYAN):    {body_fuzzy}")
    print(f"Body citations - missing (YELLOW):      {body_unmatched}")

    if fuzzy_cite_list:
        print("\n  Fuzzy matches in body:")
        for item in sorted(set(fuzzy_cite_list)):
            print(f"    ~ {item}")

    if unmatched_list:
        print("\n  Citations missing from references:")
        for item in sorted(set(unmatched_list)):
            print(f"    - {item}")

    print(f"\nReferences - cited (GREEN):             {ref_cited}")
    print(f"References - fuzzy match (CYAN):        {ref_fuzzy}")
    print(f"References - not cited (RED):           {ref_uncited}")

    if fuzzy_ref_list:
        print("\n  Fuzzy matches in references:")
        for item in sorted(set(fuzzy_ref_list)):
            print(f"    ~ {item}")

    if uncited_list:
        print("\n  Uncited references:")
        for item in sorted(set(uncited_list)):
            print(f"    - {item}")

    print(f"\n  JSON results saved for Claude Code sub-agent verification.")
    print(f"  Sub-agents will independently verify unmatched items.")
    print("\n" + "=" * 60)


if __name__ == "__main__":
    main()
